from flask import Flask, request, render_template, send_file
from flask_cors import CORS
import os
from dotenv import load_dotenv
from groq import Groq
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches

# =====================
# LOAD API KEY
# =====================
load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# =====================
# UPLOAD DIRECTORIES
# =====================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR = os.path.join(BASE_DIR, "..", "static")
UPLOAD_DIR = os.path.join(STATIC_DIR, "uploads")

LOGO_FOLDER = os.path.join(UPLOAD_DIR, "logo")
EVENT_FOLDER = os.path.join(UPLOAD_DIR, "events")
DOC_FOLDER = os.path.join(STATIC_DIR, "reports")  # For temporary docx files

os.makedirs(LOGO_FOLDER, exist_ok=True)
os.makedirs(EVENT_FOLDER, exist_ok=True)
os.makedirs(DOC_FOLDER, exist_ok=True)

# =====================
# FLASK APP
# =====================
app = Flask(
    __name__,
    template_folder="../templates",
    static_folder="../static"
)
CORS(app)

# =====================
# AI FUNCTION
# =====================
def ask_ai(prompt, mode="report"):
    try:
        if mode == "report":
            system_msg = "You write clean, professional, structured college event reports."
        elif mode == "caption":
            system_msg = "You write short, catchy, attractive social media captions."
        elif mode == "music":
            system_msg = "You suggest good songs for events or moods."
        else:
            system_msg = "You are a helpful assistant."

        res = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": prompt}
            ],
            temperature=0.6,
            max_tokens=900
        )
        return res.choices[0].message.content

    except Exception as e:
        return f"Server error: {e}"

# =====================
# ROUTES
# =====================
@app.route("/")
def home():
    return render_template("index.html")

@app.route("/chat", methods=["POST"])
def chat():
    mode = request.form.get("mode")

    # ----- MUSIC / CAPTION -----
    if mode in ["music", "caption"]:
        msg = request.form.get("message")
        text = ask_ai(msg, mode=mode)
        text = text.replace("**", "").replace("\n", "<br>")
        return text

    # ----- REPORT MODE -----
    brief = request.form.get("brief")
    logo = request.files.get("logo")
    photos = request.files.getlist("photos")

    logo_url = ""
    photo_urls = []

    # ---- Save logo ----
    if logo and logo.filename:
        filename = secure_filename(logo.filename)
        logo.save(os.path.join(LOGO_FOLDER, filename))
        logo_url = f"/static/uploads/logo/{filename}"

    # ---- Save photos ----
    for img in photos:
        if img.filename:
            filename = secure_filename(img.filename)
            img.save(os.path.join(EVENT_FOLDER, filename))
            photo_urls.append(f"/static/uploads/events/{filename}")

    # ----- AI PROMPT -----
    prompt_text = f"""
Write a professional college event report with headings:

College/School Name
Topic of Event
Date and Time
Venue
Number of Participated Students
Staff Coordinators
About the Event
Detailed Description
Additional Highlights
Event Coordinator, HOD, Principal

Based on this brief:
{brief}

Write clearly and formally.
"""
    report = ask_ai(prompt_text, mode="report").replace("**", "")

    # ----- Build HTML -----
    html = """
    <div style="font-family:Arial;max-width:900px;margin:auto;line-height:1.7;">
    """
    if logo_url:
        html += f'<div style="text-align:center;margin-bottom:20px;"><img src="{logo_url}" style="width:140px;"></div>'

    for line in report.split("\n"):
        if ":" in line:
            title, content = line.split(":", 1)
            html += f"<p><b>{title.strip()}:</b> {content.strip()}</p>"
        else:
            html += f"<p>{line}</p>"

    html += "</div>"

    # Event images grid
    if photo_urls:
        html += '<hr><h3 style="text-align:center">Event Photos</h3><div style="display:flex;flex-wrap:wrap;justify-content:center;gap:15px">'
        for img in photo_urls:
            html += f'<img src="{img}" style="width:240px;height:170px;object-fit:cover;border-radius:10px;box-shadow:0 4px 8px rgba(0,0,0,0.25)">'
        html += "</div>"

    # Add Word download button
    html += """
    <form action="/download_report" method="POST" enctype="multipart/form-data" target="_blank">
        <input type="hidden" name="brief" value='{}'>
        <input type="hidden" name="logo_path" value='{}'>
        <input type="hidden" name="photo_paths" value='{}'>
        <button type="submit" style="margin-top:20px;padding:10px 15px;font-size:16px;">Download Word Report</button>
    </form>
    """.format(brief, logo_url, ",".join(photo_urls))

    return html

# =====================
# WORD DOWNLOAD ROUTE
# =====================
@app.route("/download_report", methods=["POST"])
def download_report():
    brief = request.form.get("brief")
    logo_path = request.form.get("logo_path")
    photo_paths = request.form.get("photo_paths", "").split(",")

    # Generate docx
    doc = Document()
    doc.add_paragraph("College Event Report", style='Title')

    # Add logo
    if logo_path:
        logo_file = os.path.join(STATIC_DIR, logo_path.replace("/static/", ""))
        if os.path.exists(logo_file):
            doc.add_picture(logo_file, width=Inches(1.5))

    # Add brief / AI content
    prompt_text = f"""
Write a professional college event report with headings based on this brief:
{brief}
"""
    report = ask_ai(prompt_text, mode="report").replace("**", "")
    for line in report.split("\n"):
        doc.add_paragraph(line)

    # Add photos
    for p in photo_paths:
        if p:
            photo_file = os.path.join(STATIC_DIR, p.replace("/static/", ""))
            if os.path.exists(photo_file):
                doc.add_paragraph()  # spacing
                doc.add_picture(photo_file, width=Inches(3.5))

    # Save temp file
    file_path = os.path.join(DOC_FOLDER, "event_report.docx")
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)

# =====================
# RUN SERVER
# =====================
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5001, debug=True)

